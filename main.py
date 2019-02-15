"""
    功能:读取docx 1.提取docx所有图片并保存
    2. 将wmf图片转为png文件
    3.将docx转化为html替换html中所有图片，并保存
    4.将html转为docx p为段落，img为插入图片
    版本:v1.0
    作者:pwx322979
    时间:2018/12/15
"""
import config
import docx
import bs4
import os
from pydocx import PyDocX
import base64
import re
from docx.shared import Inches
temppath=os.getcwd()+os.path.sep+"temp"
def handle_tag(tag,parent,photo_base_path,starttext=None):
    if type(tag) == bs4.element.NavigableString:
        content = tag.replace("\n", "").strip()
        if content == "":
            pass
        else:
            run = parent.add_run()
            if tag.parent.name=="sub":
                run.font.subscript=True
            if starttext:
                content=starttext+"."+content+"\r\n"
            run.add_text(content)
    elif type(tag)==bs4.element.Tag:
        if tag.name == "p":
            p = parent.add_paragraph()
            # 遍历
            for child in tag.contents:
                handle_tag(child,p,photo_base_path)
        elif tag.name == "img":
            #按parent类型是否是进行区分
            run = parent.add_run()
            if tag["width"].endswith("pt"):
                width = float(tag["width"].replace("pt", ""))
            elif tag["width"].endswith("px"):
                width = float(tag["width"].replace("px", ""))
            if tag["height"].endswith("pt"):
                height = float(tag["height"].replace("pt", ""))
            elif tag["height"].endswith("px"):
                height = float(tag["height"].replace("px", ""))
            resultdata,path=savephoto(tag["src"], photo_base_path)
            run.add_picture(path, width=Inches(width / 80), height=Inches(height / 80))
        elif tag.name=="table":
            #获取table下所有tr标签和每一行的td标签
            tr=tag.find_all("tr")
            if len(tr)!=0:
                td = tr[0].find_all("td")
                table = parent.add_table(rows=len(tr), cols=len(td))
                startrowindex=0
                for row in tr:
                    td = row.find_all("td")
                    startcolumnindex=0
                    for column in td:
                        cell=table.cell(startrowindex,startcolumnindex)
                        p=cell.add_paragraph()
                        for child in column.contents:
                            handle_tag(child, p,photo_base_path)
                        startcolumnindex=startcolumnindex+1
                    startrowindex=startrowindex+1
        elif tag.name=="em":
            run=parent.add_run()
            if tag.parent.name=="sub":
                run.font.subscript=True
            else:
                run.font.subscript = False
            run.add_text(tag.text.replace("\n", "").strip())
        elif tag.name=="sub":
            for child in tag.contents:
                handle_tag(child,parent,photo_base_path)
        elif tag.name=="ol":
            parent = parent.add_paragraph()
            idx = 0
            if "class" in tag.attrs:
                stylename = tag.attrs["class"][0]
            for child in tag.contents:
                if child.name == "li":
                    if stylename:
                        if config.stylelist[stylename]["start"]:
                            starttext = config.stylelist[stylename]["increase"](config.stylelist[stylename]["start"],idx)
                    idx = idx + 1
                    idex=0
                    for cld in child:
                        if idex==0:
                           handle_tag(cld, parent,photo_base_path, starttext=starttext)
                        else:
                           handle_tag(cld, parent,photo_base_path)
                        idex=idex+1

        else:
            pass

def main():
    #修改为遍历时操作
    filelist=os.listdir(config.filepath)
    for file in filelist:
        if file.endswith(".docx"):
            print("start to solve ",file)
            filepath=config.filepath+"/"+file
            name=os.path.basename(filepath)
            temp_photo_path=temppath+os.path.sep+name.replace(".docx","")+os.path.sep
            if not os.path.exists(temp_photo_path):
                os.makedirs(temp_photo_path)
            # 使用pydocx转化为html
            html = PyDocX.to_html(filepath)
            bsoup = bs4.BeautifulSoup(html, "lxml")
            imglist = bsoup.find_all("img")
            for img in imglist:
                #将html中的图片保存并转换后存入html中
                img["src"],path=savephoto(img["src"], temp_photo_path)
            with open(temp_photo_path + name.replace(".docx", ".html"), "w", encoding="utf-8") as file:
                file.write(bsoup.prettify())
            #将html转为docx
            #获取所有p标签
            dstpath=temp_photo_path + name
            newdoc=docx.Document(os.getcwd()+os.path.sep+"docx/templates/default.docx")
            bodycontent = bsoup.find("body")
            bodychild = bodycontent.contents
            for child in bodychild:
                handle_tag(child, newdoc,temp_photo_path)
            newdoc.save(dstpath)
def savefile(data,path):
    with open(path, "wb") as file:
        file.write(data)
    file.close()

def savephoto(base64data,basepath):
    pattern="data:image/(\S*);base64,(\S*)"
    searchresult=re.match(pattern,base64data)
    srcpath = basepath + "temp."+searchresult[1]
    savefile(base64.b64decode(searchresult[2]), srcpath)
    if searchresult[1]=="wmf":
        dstpath = basepath + os.path.basename(srcpath).replace(".wmf", ".png")
        cmd = config.cmd + "\"" + srcpath + "\"" + " " + "\"" + dstpath + "\""
        runCmd(cmd)
        return getphotodata(dstpath),dstpath
    return getphotodata(srcpath),srcpath
def runCmd(cmd):
    p=os.popen(cmd).readlines()
    print(p)

def getphotodata(path):
    file=open(path,"rb")
    filename,type=os.path.splitext(path)
    b64_encoded_src = 'data:image/{ext};base64,{data}'.format(
        ext=type.replace(r".",""),
        data=base64.b64encode(file.read()).decode(),
    )
    file.close()
    return b64_encoded_src




if __name__=="__main__":
    main()