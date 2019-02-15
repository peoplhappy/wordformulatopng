import bs4
import docx
from PIL import Image
from docx.shared import Inches
import base64
import config
import re


path="E:/考试云平台/大学物理(上)191_03020112w.html"
dstpath="E:/1.docx"
photo_base_path="D:/pythonproject/wordimage/temp/"




def savefile(data,path):
    with open(path, "wb") as file:
        file.write(data)
    file.close()

def savephoto(base64data,basepath):
    pattern="data:image/(\S*);base64,(\S*)"
    searchresult=re.match(pattern,base64data)
    srcpath = basepath + "temp"+"."+searchresult[1]
    savefile(base64.b64decode(searchresult[2]), srcpath)
    return searchresult[1]

def handle_tag(tag,parent,starttext=None):
    print("tag:",tag.name)
    if type(tag) == bs4.element.NavigableString:
        content=tag.replace("\n", "").strip()
        if content=="":
            pass
        else:
            run=parent.add_run()
            if tag.parent.name=="sub":
                run.font.subscript=True
            if starttext:
                content=starttext+"."+content+"\r\n"
            run.add_text(content)
    elif type(tag)==bs4.element.Tag:
        if tag.name == "p":
            p = parent.add_paragraph()
            # 遍历
            for child in tag.contents:
                handle_tag(child,p)
        elif tag.name == "img":
            #按parent类型是否是进行区分
            run = parent.add_run()
            if tag["width"].endswith("pt"):
                width = float(tag["width"].replace("pt", ""))
            elif tag["width"].endswith("px"):
                width = float(tag["width"].replace("px", ""))
            if tag["height"].endswith("pt"):
                height = float(tag["height"].replace("pt", ""))
            elif tag["height"].endswith("px"):
                height = float(tag["height"].replace("px", ""))
            afterfix=savephoto(tag["src"], photo_base_path)
            run.add_picture(photo_base_path + "temp"+"."+afterfix, width=Inches(width / 80), height=Inches(height / 80))
        elif tag.name=="table":
            #获取table下所有tr标签和每一行的td标签
            tr=tag.find_all("tr")
            if len(tr)!=0:
                td = tr[0].find_all("td")
                table = parent.add_table(rows=len(tr), cols=len(td))
                startrowindex=0
                for row in tr:
                    td = row.find_all("td")
                    startcolumnindex=0
                    for column in td:
                        cell=table.cell(startrowindex,startcolumnindex)
                        p=cell.add_paragraph()
                        for child in column.contents:
                            handle_tag(child, p)
                        startcolumnindex=startcolumnindex+1
                    startrowindex=startrowindex+1
        elif tag.name=="em":
            run=parent.add_run()
            if tag.parent.name=="sub":
                run.font.subscript=True
            else:
                run.font.subscript = False
            run.add_text(tag.string.replace("\n", "").strip())
        elif tag.name=="sub":
            for child in tag.contents:
                handle_tag(child,parent)
        elif tag.name=="ol":
            parent=parent.add_paragraph()
            idx=0
            if "class" in tag.attrs:
                stylename=tag.attrs["class"][0]
            for child in tag.contents:
                if child.name=="li":
                    if stylename:
                       if config.stylelist[stylename]["start"]:
                           starttext=config.stylelist[stylename]["increase"](config.stylelist[stylename]["start"],idx)
                           idx=idx+1
                    for cld in child:
                        handle_tag(cld,parent,starttext)

        else:
            pass

if __name__=="__main__":
    newdoc = docx.Document()
    bsoup = bs4.BeautifulSoup(open(path, "r", encoding="utf-8"), "lxml")
    bodycontent = bsoup.find("body")
    bodychild = bodycontent.contents
    for child in bodychild:
        handle_tag(child, newdoc)
    newdoc.save(dstpath)